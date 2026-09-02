"""
Generador de memoria de cálculo (.docx), con la misma estructura que las
memorias de referencia del proyecto (anexo MEMORIAS_DE_CALCULO_1.docx):
portada, introducción, evaluación de cargas, materiales, secciones,
geometría, combinaciones de carga, sistema estructural, datos de entrada
y esfuerzos/verificación final — de forma que cualquier calculista
familiarizado con estas memorias reconozca inmediatamente el formato.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from ..geometry.model import RackModel, Section
from ..loads.combinations import Combination
from ..loads.seismic import SeismicResult


@dataclass
class ProjectInfo:
    titulo: str
    ciudad: str
    fecha: str
    ingeniero: str
    especialidad: str
    matricula: str
    cliente: str = ""
    descripcion: str = ""
    software_version: str = "Vortex 0.1.0"


@dataclass
class ReportData:
    project: ProjectInfo
    model: RackModel
    dl_note: str
    ll_kn_m2: float
    pl_per_level_kn: float
    seismic_transversal: Optional[SeismicResult]
    seismic_longitudinal: Optional[SeismicResult]
    material_names: List[str]
    upright_section: Section
    beam_section: Section
    brace_section: Section
    method_name: str
    combos: List[Combination]
    design_rows: List[Dict[str, Any]] = field(default_factory=list)
    # Filas completas de vortex.analysis.pipeline.MemberResultRow (opcional):
    # si se suministran, se agregan las tablas "RESISTENCIA <sección> MODELO
    # CFS" y "CHEQUEO" con el mismo formato de columnas (H, P, Mx, Vy, My,
    # Vx) que usa el calculista de referencia en su memoria (ver
    # `_add_resistencia_chequeo_tables`). Se accede por duck-typing
    # (.kind, .length_m, .upright_check) para no acoplar `report` a
    # `analysis`.
    member_rows_detail: List[Any] = field(default_factory=list)
    disclaimer_extra: str = ""
    # Texto de las recomendaciones del asesor de IA (Groq) ya consultado
    # desde la GUI (ver `vortex.ai.advisor` / `vortex.gui.app`), para
    # incluirlo como sección de la memoria (numeral 7). Vacío si el
    # usuario no consultó la IA antes de exportar — en ese caso la sección
    # se genera igual, con una nota indicándolo (ver `_add_ai_analysis`).
    ai_analysis: str = ""


def _add_title_page(doc: Document, data: ReportData) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MEMORIA DE CÁLCULO")
    run.bold = True
    run.font.size = Pt(20)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(data.project.titulo).font.size = Pt(14)

    doc.add_paragraph()
    info = [
        ("Contiene", data.project.titulo),
        ("Cliente", data.project.cliente),
        ("Ciudad", data.project.ciudad),
        ("Fecha", data.project.fecha),
    ]
    for label, value in info:
        if not value:
            continue
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run(f"{label}: {value}")

    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("_____________________________")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Firma: {data.project.ingeniero}").bold = True
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(data.project.especialidad)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run(f"M.P. {data.project.matricula}")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        f"Generada con {data.project.software_version} — documento sujeto a "
        f"revisión, ajuste y aprobación por el ingeniero calculista responsable "
        f"antes de su uso para construcción."
    )
    r.italic = True
    r.font.size = Pt(9)
    doc.add_page_break()


def _add_toc(doc: Document) -> None:
    doc.add_heading("TABLA DE CONTENIDO", level=1)
    items = [
        "Introducción",
        "1. Procedimiento estructural",
        "1.1 Evaluación de cargas",
        "1.2 Materiales",
        "1.3 Secciones",
        "1.4 Geometría",
        "2. Combinaciones de carga",
        "3. Definición del sistema estructural",
        "4. Datos de entrada",
        "5. Esfuerzos y verificación de elementos",
        "6. Índice de elementos (referencia rápida)",
        "7. Análisis asistido por IA",
        "8. Conclusiones",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_page_break()


def _add_intro(doc: Document, data: ReportData) -> None:
    doc.add_heading("INTRODUCCIÓN", level=1)
    doc.add_paragraph(
        f"El resultado del estudio que a continuación se presenta consiste en el "
        f"análisis matemático y diseño estructural correspondiente a: "
        f"{data.project.titulo}."
    )
    doc.add_paragraph(
        "En esta memoria se define el procedimiento utilizado para el análisis "
        "y diseño estructural: el método de análisis empleado (elementos finitos "
        "de pórtico espacial, método de la rigidez directa, con conexiones "
        "viga-paral semirrígidas), el sistema estructural y los resultados y "
        "conclusiones obtenidos."
    )
    doc.add_paragraph(
        "Todo el análisis y diseño estructural se realizó cumpliendo con las "
        "exigencias de la norma NTC 5689 \"Especificación para el diseño, "
        "ensayo y utilización de estanterías industriales de acero\" (adopción "
        "modificada de ANSI/RMI MH16.1), con apoyo de AISI \"Specification for "
        "the Design of Cold-Formed Steel Structural Members\" y AISC "
        "\"Specification for Structural Steel Buildings\"."
    )


def _add_loads_section(doc: Document, data: ReportData) -> None:
    doc.add_heading("1. PROCEDIMIENTO ESTRUCTURAL", level=1)
    doc.add_paragraph(
        "La estantería se debe chequear para que tenga la resistencia y "
        "rigidez adecuadas ante las solicitaciones mínimas exigidas en la "
        "norma NTC 5689."
    )
    doc.add_heading("1.1 Evaluación de cargas", level=2)

    doc.add_heading("1.1.1 Cargas muertas (DL)", level=3)
    doc.add_paragraph(data.dl_note or "DL = Peso propio de la estructura (calculado por elemento).")

    doc.add_heading("1.1.2 Cargas vivas (LL) y de producto (PL)", level=3)
    doc.add_paragraph(f"LL = {data.ll_kn_m2:.2f} kN/m² (cargas no asociadas a estibas/producto).")
    doc.add_paragraph(
        f"PL = {data.pl_per_level_kn:.2f} kN por nivel y por bahía (carga máxima "
        f"de estibas/producto almacenado, según placa de capacidad — numeral 1.5.2)."
    )

    doc.add_heading("1.1.3 Cargas sísmicas (EL) — NTC 5689 numeral 2.7", level=3)
    for label, res in (
        ("Dirección transversal (marcos, arriostrada)", data.seismic_transversal),
        ("Dirección longitudinal (vigas, no arriostrada)", data.seismic_longitudinal),
    ):
        if res is None:
            continue
        doc.add_paragraph(label, style="Intense Quote")
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        rows = [
            ("Tipo de perfil de suelo", res.soil_type),
            ("Aa", f"{res.aa:.3f}"),
            ("Av", f"{res.av:.3f}"),
            ("Ca (Tabla 1, NTC 5689)", f"{res.ca:.4f}"),
            ("Cv (Tabla 2, NTC 5689)", f"{res.cv:.4f}"),
            ("R", f"{res.r:.2f}"),
            ("Ip", f"{res.ip:.2f}"),
            ("PLRF", f"{res.plrf:.3f}"),
            ("Ws = 0.67·PLRF·PL + DL + 0.25·LL", f"{res.ws:.2f} kN"),
            ("T (periodo fundamental)", f"{res.period_s:.3f} s" if res.period_s else "no calculado (método simplificado)"),
            ("Cs = Cs de diseño", f"{res.cs:.5f}"),
            ("V = Cs·Ip·Ws (cortante sísmico de base)", f"{res.v_base:.2f} kN"),
        ]
        for label_r, val_r in rows:
            r = table.add_row().cells
            r[0].text, r[1].text = label_r, val_r
        doc.add_paragraph()


def _add_materials_sections(doc: Document, data: ReportData) -> None:
    doc.add_heading("1.2 Materiales", level=2)
    doc.add_paragraph("Propiedades de diseño asumidas:")
    for name in data.material_names:
        doc.add_paragraph(f"• {name}", style="List Bullet")

    doc.add_heading("1.3 Secciones", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Elemento", "Sección", "A (cm²)", "Iy (cm⁴)", "Iz (cm⁴)", "Fy (MPa)"]):
        hdr[i].text = h
    for label, sec in (
        ("Paral", data.upright_section),
        ("Viga", data.beam_section),
        ("Diagonal", data.brace_section),
    ):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = sec.name
        row[2].text = f"{sec.A * 1e4:.2f}"
        row[3].text = f"{sec.Iy * 1e8:.2f}"
        row[4].text = f"{sec.Iz * 1e8:.2f}"
        row[5].text = f"{sec.Fy / 1000.0:.0f}"

    doc.add_heading("1.4 Geometría", level=2)
    m = data.model
    doc.add_paragraph(f"Altura total: {m.level_elevations[-1]:.2f} m" if m.level_elevations else "")
    doc.add_paragraph(f"Número de niveles de carga: {m.n_levels}")
    doc.add_paragraph(f"Número de bahías: {m.n_bays}")
    doc.add_paragraph(f"Longitud de viga: {m.bay_length:.2f} m")
    doc.add_paragraph(f"Profundidad de marco: {m.frame_depth:.2f} m")
    if m.level_heights:
        doc.add_paragraph("Alturas libres entre niveles:")
        for i, h in enumerate(m.level_heights):
            doc.add_paragraph(f"De nivel {i} a {i + 1} = {h:.2f} m", style="List Bullet")


def _add_combinations(doc: Document, data: ReportData) -> None:
    doc.add_heading("2. COMBINACIONES DE CARGA", level=1)
    doc.add_paragraph(
        f"Método de diseño: {data.method_name} — NTC 5689 numeral "
        f"{'2.1' if data.method_name == 'ASD' else '2.2'}."
    )
    doc.add_paragraph("DEFINICIONES:")
    defs = [
        ("DL", "Carga muerta"),
        ("LL", "Carga viva distinta a la de estibas/producto"),
        ("PL", "Máxima carga de estibas o productos almacenados"),
        ("PL_APP", "Carga de producto aplicable para arrancamiento"),
        ("SL / RL", "Carga de granizo / lluvia"),
        ("WL", "Carga de viento"),
        ("EL", "Carga sísmica"),
        ("IMP", "Carga de impacto vertical en un entrepaño"),
    ]
    for k, v in defs:
        doc.add_paragraph(f"{k} = {v}", style="List Bullet")

    doc.add_paragraph("COMBINACIONES:")
    for c in data.combos:
        doc.add_paragraph(f"{c.label()}  —  {c.description} ({c.scope.value})")


def _add_structural_system(doc: Document, data: ReportData) -> None:
    doc.add_heading("3. DEFINICIÓN DEL SISTEMA ESTRUCTURAL", level=1)
    doc.add_paragraph(
        "El modelo matemático de la estantería se elaboró a partir de los "
        "elementos que la conforman: parales, vigas y diagonales de "
        "arriostramiento, mediante análisis matricial de pórtico espacial "
        "(3D, 6 grados de libertad por nudo, método de la rigidez directa)."
    )
    doc.add_paragraph(
        "Las vigas porta-estibas se conectan a los parales mediante "
        "conexiones semirrígidas (numeral 7.1), modeladas como resortes "
        "rotacionales de rigidez km obtenida del ensayo tipo cantiléver "
        "(numeral 9.4.1) o de un valor de referencia cuando no se dispone "
        "de dicho ensayo (ver notas de advertencia en los resultados)."
    )
    doc.add_paragraph(
        "El marco transversal (plano paral frente-paral fondo) se "
        "arriostra con diagonales tipo armadura (extremos articulados), "
        "R=4 según numeral 2.7.3. La dirección longitudinal (a lo largo "
        "del corredor) es un pórtico no arriostrado que resiste cargas "
        "laterales por flexión de parales y vigas a través de las "
        "conexiones semirrígidas, R=6."
    )


def _add_input_data(doc: Document, data: ReportData) -> None:
    doc.add_heading("4. DATOS DE ENTRADA", level=1)
    m = data.model
    doc.add_paragraph(f"Table: Node Coordinates  (total nudos: {len(m.nodes)})")
    doc.add_paragraph(f"Table: Frame Section Assignments  (total elementos: {len(m.members)})")
    doc.add_paragraph(
        f"Table: Frame Loads - Distributed  (DL por elemento, PL={data.pl_per_level_kn:.2f} "
        f"kN/nivel-bahía repartido en vigas)"
    )
    doc.add_paragraph("Table: Joint Loads - Force  (cargas sísmicas Fx por nivel, ver numeral 1.1.3)")


def _add_results(doc: Document, data: ReportData) -> None:
    doc.add_heading("5. ESFUERZOS Y VERIFICACIÓN DE ELEMENTOS", level=1)
    if not data.design_rows:
        doc.add_paragraph("(sin resultados de verificación cargados en este reporte)")
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Elemento", "Tipo", "Combinación crítica", "Demanda/Capacidad", "Ratio", "Estado"]):
        hdr[i].text = h
    n_fail = 0
    for row in data.design_rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get("elemento", ""))
        cells[1].text = str(row.get("tipo", ""))
        cells[2].text = str(row.get("combo", ""))
        cells[3].text = str(row.get("demanda_capacidad", ""))
        ratio = row.get("ratio", 0.0)
        cells[4].text = f"{ratio:.2f}"
        estado = "OK" if ratio <= 1.0 else "NO CUMPLE"
        if ratio > 1.0:
            n_fail += 1
        cells[5].text = estado

    doc.add_paragraph()
    if n_fail == 0:
        doc.add_paragraph(
            f"Todos los elementos verificados ({len(data.design_rows)}) cumplen "
            f"con una relación demanda/capacidad ≤ 1.0 para la combinación "
            f"crítica reportada."
        )
    else:
        p = doc.add_paragraph()
        p.add_run(
            f"ATENCIÓN: {n_fail} de {len(data.design_rows)} elementos NO cumplen "
            f"(relación demanda/capacidad > 1.0). Revisar sección, geometría o "
            f"cargas antes de continuar."
        ).bold = True


def _add_resistencia_chequeo_tables(doc: Document, data: ReportData) -> None:
    """
    Tablas "RESISTENCIA <sección> MODELO CFS" y "CHEQUEO", con el mismo
    nombre, orden de columnas y convención de ejes (Mx=M3/V2=Vy en el
    plano 1, My=M2/V3=Vx en el plano 2) que usa el calculista de
    referencia en su memoria de cálculo — para que el formato sea
    inmediatamente reconocible frente a memorias anteriores del mismo
    proyecto, aunque los valores sean los que calcula Vortex.
    """
    rows = [r for r in data.member_rows_detail if getattr(r, "kind", "") == "Paral" and r.upright_check]
    if not rows:
        return

    doc.add_heading(f"RESISTENCIA {data.upright_section.name} MODELO CFS", level=2)
    doc.add_paragraph(
        "Fuerzas de diseño en la combinación crítica de cada paral (ver "
        "numeral 2), en la convención de ejes del anexo: Mx/Vy = momento/"
        "cortante en el plano 1 (M3/V2 de Vortex), My/Vx = momento/"
        "cortante en el plano 2 (M2/V3 de Vortex)."
    )
    t1 = doc.add_table(rows=1, cols=7)
    t1.style = "Light Grid Accent 1"
    for i, h in enumerate(["ITEM", "H [m]", "P [KN]", "Mx [KN-m]", "Vy [KN]", "My [KN-m]", "Vx [KN]"]):
        t1.rows[0].cells[i].text = h
    for item, row in enumerate(rows, start=1):
        c = row.upright_check
        cells = t1.add_row().cells
        cells[0].text = str(item)
        cells[1].text = f"{row.length_m:.2f}"
        cells[2].text = f"{c.P:.3f}"
        cells[3].text = f"{c.M3:.4f}"
        cells[4].text = f"{c.V2:.3f}"
        cells[5].text = f"{c.M2:.4f}"
        cells[6].text = f"{c.V3:.3f}"

    doc.add_heading("CHEQUEO", level=2)
    doc.add_paragraph(
        "Las columnas P, Mx, Vy, My y Vx comparan cada componente de forma "
        "INDEPENDIENTE contra su propia capacidad admisible (Pa, Ma3, Va, "
        "Ma2, Va) — son un diagnóstico rápido de qué acción es más "
        "exigente, pero NO son el chequeo de diseño real de un elemento a "
        "flexo-compresión: un paral puede tener sus 5 componentes por "
        "separado por debajo de 1.0 y aun así NO CUMPLIR, porque la "
        "interacción P/Pa + M2/Ma2 + M3/Ma3 (con amplificación por efectos "
        "de segundo orden) puede superar 1.0 aunque ningún término lo haga "
        "por sí solo. La columna \"Interacción\" es la que gobierna el "
        "resultado ESTADO de la sección 5 de esta memoria; si difiere de "
        "\"OK\" en todos los componentes individuales, es precisamente por "
        "ese efecto de interacción y NO es una inconsistencia del cálculo."
    )
    t2 = doc.add_table(rows=1, cols=8)
    t2.style = "Light Grid Accent 1"
    headers2 = ["ITEM", "P [KN]", "Mx [KN-m]", "Vy [KN]", "My [KN-m]", "Vx [KN]", "Interacción", "ESTADO"]
    for i, h in enumerate(headers2):
        t2.rows[0].cells[i].text = h
    n_fail_component = 0
    n_fail_interaction = 0
    for item, row in enumerate(rows, start=1):
        c = row.upright_check
        checks = c.component_checks
        vals = [checks["P"], checks["M3"], checks["V2"], checks["M2"], checks["V3"]]
        if not all(vals):
            n_fail_component += 1
        governing_fail = c.ratio_interaction > 1.0
        if governing_fail:
            n_fail_interaction += 1
        cells = t2.add_row().cells
        cells[0].text = str(item)
        for i, ok in enumerate(vals, start=1):
            cells[i].text = "OK" if ok else "NO CUMPLE"
        cells[6].text = f"{c.ratio_interaction:.2f}"
        cells[7].text = "NO CUMPLE" if governing_fail else "OK"
    doc.add_paragraph()
    if n_fail_interaction:
        p = doc.add_paragraph()
        p.add_run(
            f"{n_fail_interaction} de {len(rows)} parales NO CUMPLEN la interacción "
            f"P/Pa + M2/Ma2 + M3/Ma3 (columna \"Interacción\", ratio > 1.0) — este es "
            f"el resultado gobernante, coherente con la sección 5. Además, "
            f"{n_fail_component} de {len(rows)} tienen al menos un componente "
            f"individual (P, Mx, Vy, My o Vx) que por sí solo ya excede su "
            f"capacidad admisible."
        ).bold = True
    else:
        doc.add_paragraph(
            f"Los {len(rows)} parales cumplen tanto en la interacción P/Pa+M2/Ma2+"
            f"M3/Ma3 como en cada componente individual."
        )


def _add_element_index(doc: Document, data: ReportData) -> None:
    """
    Índice de TODOS los elementos del modelo (etiqueta -> ubicación física:
    marco/bahía, lado, nivel y nudos de extremo con coordenadas), para que
    el calculista pueda ubicar de inmediato, ante cualquier duda, a qué
    elemento real corresponde una etiqueta mencionada en las secciones 5 o
    "CHEQUEO" (p.ej. "VIGA B0-F N3" o "PARAL M1-P N2-N3"). "F"/"P" en la
    etiqueta identifican el lado del marco: F=frente (Y=0), P=fondo/
    posterior (Y=profundidad de marco) — antes ambos lados usaban la letra
    "F" (frente y fondo empiezan igual en español) y dos elementos físicos
    distintos terminaban con el mismo nombre en la memoria.
    """
    doc.add_heading("6. ÍNDICE DE ELEMENTOS (REFERENCIA RÁPIDA)", level=1)
    doc.add_paragraph(
        "F = lado frente del marco (Y=0). P = lado fondo/posterior del "
        "marco (Y=profundidad de marco). Coordenadas de nudo en metros, "
        "origen en la esquina frente-izquierda del piso."
    )
    m = data.model
    members_sorted = sorted(
        m.members.values(),
        key=lambda mm: (mm.kind.name, mm.frame_index or 0, mm.bay_index or 0, mm.level_index or 0, mm.side or ""),
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Etiqueta", "Tipo", "Marco/Bahía", "Lado", "Nivel", "Nudo i (x,y,z)", "Nudo j (x,y,z)"]):
        table.rows[0].cells[i].text = h
    for mm in members_sorted:
        ni, nj = m.nodes[mm.node_i], m.nodes[mm.node_j]
        row = table.add_row().cells
        row[0].text = mm.label
        row[1].text = mm.kind.name.capitalize()
        row[2].text = (
            f"Marco {mm.frame_index}" if mm.bay_index is None else f"Bahía {mm.bay_index}"
        ) if (mm.frame_index is not None or mm.bay_index is not None) else "—"
        row[3].text = {"frente": "F (frente)", "fondo": "P (fondo)"}.get(mm.side or "", "—")
        row[4].text = str(mm.level_index) if mm.level_index is not None else "—"
        row[5].text = f"({ni.x:.2f}, {ni.y:.2f}, {ni.z:.2f})"
        row[6].text = f"({nj.x:.2f}, {nj.y:.2f}, {nj.z:.2f})"


def _looks_like_real_ai_text(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip()
    if not stripped:
        return False
    placeholders = ("consultando ia", "⚠")
    low = stripped.lower()
    return not any(low.startswith(p) for p in placeholders)


def _add_ai_analysis(doc: Document, data: ReportData) -> None:
    doc.add_heading("7. ANÁLISIS ASISTIDO POR IA", level=1)
    if _looks_like_real_ai_text(data.ai_analysis):
        doc.add_paragraph(
            "El siguiente análisis fue generado por un modelo de lenguaje "
            "(LLM, vía la API de Groq) a partir del resumen numérico de "
            "este mismo cálculo (sismo, elementos más críticos, relación "
            "demanda/capacidad). Es una ayuda de revisión adicional, NO "
            "un chequeo normativo ni un reemplazo del criterio del "
            "ingeniero calculista: puede pasar por alto errores o "
            "malinterpretar el contexto del proyecto real."
        ).italic = True
        for line in data.ai_analysis.strip().splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith(("-", "•", "*")):
                doc.add_paragraph(line.lstrip("-•* ").strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph(
            "No se generaron recomendaciones de IA para este reporte "
            "(use el botón \"Recomendaciones IA\" en Vortex antes de "
            "exportar la memoria para incluir aquí ese análisis)."
        )


def _add_conclusions(doc: Document, data: ReportData) -> None:
    doc.add_heading("8. CONCLUSIONES", level=1)
    doc.add_paragraph(
        "Los elementos de la estantería descrita en esta memoria fueron "
        "verificados bajo las combinaciones de carga de la norma NTC 5689, "
        "incluyendo la solicitación sísmica calculada según el numeral 2.7."
    )
    doc.add_paragraph(
        "Este documento fue generado automáticamente por un software de "
        "apoyo al cálculo (Vortex) y DEBE ser revisado, complementado y "
        "firmado por el ingeniero calculista responsable antes de su uso "
        "para fabricación o construcción. En particular, deben verificarse: "
        "las propiedades certificadas de las secciones (incluyendo Cw, xo, "
        "ro para pandeo flexo-torsional), la rigidez de conexión real "
        "obtenida de ensayo (numeral 9.4), y la capacidad de anclajes al "
        "concreto según ACI 318 capítulo 17 con la geometría real del "
        "proyecto."
    )
    if data.disclaimer_extra:
        doc.add_paragraph(data.disclaimer_extra)


def generate_memoria(data: ReportData, output_path: str) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(doc, data)
    _add_toc(doc)
    _add_intro(doc, data)
    _add_loads_section(doc, data)
    _add_materials_sections(doc, data)
    _add_combinations(doc, data)
    _add_structural_system(doc, data)
    _add_input_data(doc, data)
    _add_results(doc, data)
    _add_resistencia_chequeo_tables(doc, data)
    _add_element_index(doc, data)
    _add_ai_analysis(doc, data)
    _add_conclusions(doc, data)

    doc.save(output_path)
    return output_path
