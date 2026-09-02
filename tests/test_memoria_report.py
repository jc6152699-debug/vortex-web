"""
Contenido de la memoria de cálculo (.docx) exportada:
  - la tabla "CHEQUEO" debe mostrar la interacción P/Pa+M2/Ma2+M3/Ma3 y el
    ESTADO gobernante (no sólo el chequeo independiente por componente, que
    por sí solo puede decir "OK" en los 5 componentes mientras el elemento
    en realidad NO CUMPLE por interacción -- ver `design.upright_cfs`);
  - el índice de elementos permite ubicar cualquier etiqueta mencionada en
    los resultados (p.ej. "VIGA B0-F N3") por marco/bahía, lado y nudos;
  - la sección de IA se incluye siempre, con el texto real cuando se
    suministra o una nota explícita cuando no.
"""
from docx import Document

from vortex.geometry import RackParameters, build_selective_rack
from vortex.sections.catalog import default_catalog
from vortex.analysis import PipelineInputs, SeismicInputs, run_full_check
from vortex.report import ProjectInfo, ReportData, generate_memoria
from vortex.units import kgf_to_kn


def _build_report_data(ai_analysis: str = "") -> ReportData:
    catalog = default_catalog()
    params = RackParameters(
        n_bays=2, bay_length=2.44, frame_depth=1.06, level_heights=[1.20, 1.80, 1.80],
        upright_section=catalog["PARAL 122x2.5mm"],
        beam_section=catalog["VIGA CAJA 160x60x1.5mm"],
        brace_section=catalog["DIAGONAL TUBULAR 30x30x2.0mm"],
        base_fixity="pinned",
    )
    model = build_selective_rack(params)
    inputs = PipelineInputs(
        # Carga de producto deliberadamente alta para asegurar que al
        # menos un paral NO CUMPLA en este modelo pequeño de prueba.
        pl_per_level_kn=kgf_to_kn(12000.0), ll_kn_m2=0.0,
        seismic=SeismicInputs(soil_type="D", aa=0.15, av=0.20),
    )
    result = run_full_check(model, inputs)
    design_rows = [
        {"elemento": r.label, "tipo": r.kind, "combo": r.combo,
         "demanda_capacidad": r.detail, "ratio": r.ratio}
        for r in result.member_rows.values()
    ]
    project = ProjectInfo(
        titulo="TEST", ciudad="Medellín", fecha="2026-09-01",
        ingeniero="X", especialidad="Y", matricula="Z",
    )
    return ReportData(
        project=project, model=model, dl_note="n", ll_kn_m2=0.0,
        pl_per_level_kn=inputs.pl_per_level_kn,
        seismic_transversal=result.seismic_transversal,
        seismic_longitudinal=result.seismic_longitudinal,
        material_names=[catalog["PARAL 122x2.5mm"].material.name],
        upright_section=catalog["PARAL 122x2.5mm"],
        beam_section=catalog["VIGA CAJA 160x60x1.5mm"],
        brace_section=catalog["DIAGONAL TUBULAR 30x30x2.0mm"],
        method_name="LRFD", combos=result.combos, design_rows=design_rows,
        member_rows_detail=list(result.member_rows.values()),
        ai_analysis=ai_analysis,
    )


def _table_after_heading(doc: Document, heading_text: str):
    # docx no expone directamente "la tabla después del párrafo N" por
    # índice de párrafo/tabla lógico; se recorre el XML del body en orden
    # y se toma la primera <w:tbl> que aparece después del <w:p> que
    # contiene el texto del encabezado buscado.
    from docx.table import Table
    body_children = list(doc.element.body)
    seen_heading = False
    for child in body_children:
        if not seen_heading:
            if child.tag.endswith("}p") and heading_text in "".join(
                n.text or "" for n in child.iter() if n.tag.endswith("}t")
            ):
                seen_heading = True
            continue
        if child.tag.endswith("}tbl"):
            return Table(child, doc)
    raise AssertionError(f"no se encontró una tabla después de {heading_text!r}")


def test_chequeo_table_shows_interaction_and_estado_columns(tmp_path):
    data = _build_report_data()
    out = tmp_path / "m.docx"
    generate_memoria(data, str(out))
    doc = Document(str(out))

    table = _table_after_heading(doc, "CHEQUEO")
    header = [c.text for c in table.rows[0].cells]
    assert "Interacción" in header
    assert "ESTADO" in header

    # Al menos una fila con ESTADO = NO CUMPLE debe tener Interacción > 1.0
    idx_estado = header.index("ESTADO")
    idx_inter = header.index("Interacción")
    any_checked = False
    for row in table.rows[1:]:
        estado = row.cells[idx_estado].text
        inter = float(row.cells[idx_inter].text)
        if estado == "NO CUMPLE":
            assert inter > 1.0
            any_checked = True
        else:
            assert inter <= 1.0
    assert any_checked, "se esperaba al menos un paral NO CUMPLE en este modelo de prueba"


def test_element_index_lists_frente_and_fondo_separately(tmp_path):
    data = _build_report_data()
    out = tmp_path / "m.docx"
    generate_memoria(data, str(out))
    doc = Document(str(out))

    table = _table_after_heading(doc, "ÍNDICE DE ELEMENTOS")
    labels = [row.cells[0].text for row in table.rows[1:]]
    assert len(labels) == len(set(labels)), "hay etiquetas repetidas en el índice de elementos"
    assert any("VIGA B0-F" in lbl for lbl in labels)
    assert any("VIGA B0-P" in lbl for lbl in labels)


def test_ai_section_shows_placeholder_when_empty(tmp_path):
    data = _build_report_data(ai_analysis="")
    out = tmp_path / "m.docx"
    generate_memoria(data, str(out))
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "No se generaron recomendaciones de IA" in full_text


def test_ai_section_includes_real_analysis_text(tmp_path):
    data = _build_report_data(ai_analysis="- Revisar el paral M1-F N0-N1 por alta utilización.")
    out = tmp_path / "m.docx"
    generate_memoria(data, str(out))
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Revisar el paral M1-F N0-N1 por alta utilización." in full_text
